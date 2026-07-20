"""
reader.py
=========
JOB OF THIS FILE: take an uploaded file and turn it into plain text.

We keep track of WHERE each piece of text came from (page number for PDFs,
paragraph number for Word). That "where" is what lets us later tell the user
"this answer came from page 4" instead of just trusting the AI.

Output shape (same for both file types) -> a list of dicts:
    [{"page": 1, "text": "....."}, {"page": 2, "text": "....."}]
"""

import io

from pypdf import PdfReader
from docx import Document


# ---------------------------------------------------------------------------
# STEP 1: PDF reading
# ---------------------------------------------------------------------------
def read_pdf(file_bytes: bytes) -> list[dict]:
    """Pull the text out of a PDF, one entry per page."""

    # PdfReader wants a file-like object, so we wrap the raw bytes in BytesIO.
    reader = PdfReader(io.BytesIO(file_bytes))

    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        # extract_text() can return None on image-only pages -> default to "".
        text = page.extract_text() or ""
        text = _clean(text)

        # Skip pages that are basically empty (scanned images, blank pages).
        if text:
            pages.append({"page": page_number, "text": text})

    return pages


# ---------------------------------------------------------------------------
# STEP 2: Word reading
# ---------------------------------------------------------------------------
def read_docx(file_bytes: bytes) -> list[dict]:
    """
    Pull text out of a .docx file.

    Word has no real "pages" (page breaks depend on the printer/screen), so we
    fake it: every ~25 paragraphs becomes one "page". This keeps the citation
    numbers small and useful instead of saying "paragraph 412".
    """
    document = Document(io.BytesIO(file_bytes))

    # 2a. Grab normal paragraphs.
    blocks = [p.text for p in document.paragraphs if p.text.strip()]

    # 2b. Grab table text too - lots of real documents hide data in tables.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    # 2c. Group the paragraphs into fake pages.
    pages = []
    group_size = 25
    for index in range(0, len(blocks), group_size):
        chunk_of_paragraphs = blocks[index:index + group_size]
        text = _clean("\n".join(chunk_of_paragraphs))
        if text:
            pages.append({"page": (index // group_size) + 1, "text": text})

    return pages


# ---------------------------------------------------------------------------
# STEP 3: One entry point the rest of the app calls
# ---------------------------------------------------------------------------
def read_file(filename: str, file_bytes: bytes) -> list[dict]:
    """Decide which reader to use based on the file extension."""
    name = filename.lower()

    if name.endswith(".pdf"):
        return read_pdf(file_bytes)

    if name.endswith(".docx"):
        return read_docx(file_bytes)

    # .doc (old Word format) is a completely different binary format that
    # python-docx cannot open. Tell the user clearly instead of failing weirdly.
    if name.endswith(".doc"):
        raise ValueError("Old .doc files aren't supported. Save it as .docx and upload again.")

    raise ValueError("Only PDF and Word (.docx) files are supported.")


# ---------------------------------------------------------------------------
# Small helper: tidy up messy extracted text
# ---------------------------------------------------------------------------
def _clean(text: str) -> str:
    """Remove weird spacing that PDF extraction usually leaves behind."""
    # Replace non-breaking spaces and collapse runs of blank lines.
    text = text.replace("\xa0", " ")
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines).strip()
